"""
Generate build_plan.docx — full end-to-end build plan for NativeReady.
Every step from zero to launched website, with risks and mitigations.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)


def add_heading(text, level=1):
    return doc.add_heading(text, level=level)


def add_paragraph(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_meta(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x6a, 0x6a, 0x73)


def add_kv(key, value):
    p = doc.add_paragraph()
    p.add_run(f'{key}. ').bold = True
    p.add_run(value)


def add_phase(num, title, goal, steps, deliverables, time, risks):
    add_heading(f'Phase {num} — {title}', level=2)
    add_kv('Goal', goal)
    add_kv('Time estimate', time)

    p = doc.add_paragraph()
    p.add_run('Steps').bold = True
    for s in steps:
        doc.add_paragraph(s, style='List Number')

    p = doc.add_paragraph()
    p.add_run('Deliverables').bold = True
    for d in deliverables:
        doc.add_paragraph(d, style='List Bullet')

    p = doc.add_paragraph()
    p.add_run('Risks and mitigations').bold = True
    for risk, mitigation in risks:
        rp = doc.add_paragraph()
        rp.add_run('Risk: ').bold = True
        rp.add_run(risk)
        mp = doc.add_paragraph()
        mp.add_run('  Mitigation: ').italic = True
        mp.add_run(mitigation).italic = True

    doc.add_paragraph()


# ============================================================
# TITLE
# ============================================================
title = doc.add_heading('NativeReady — End-to-End Build Plan', 0)

subtitle = doc.add_paragraph()
sub = subtitle.add_run(
    'From zero to launched website. Every step, every risk, every mitigation.'
)
sub.italic = True
sub.font.size = Pt(13)

add_meta('Build plan  ·  NativeReady project  ·  Brhanu Fentaw, April 2026')

doc.add_paragraph()

# ============================================================
# EXECUTIVE SUMMARY
# ============================================================
add_heading('Executive summary', level=1)

doc.add_paragraph(
    'NativeReady is a web tool that takes a protein sequence and predicts '
    'whether it will give usable native mass spec data, before the experiment '
    'is run. This document specifies every step from a blank folder to a '
    'public, payment-enabled, fully functional website.'
)

add_kv('Total time estimate', '8 to 12 weekends (60 to 100 hours of focused work)')
add_kv('Total cost estimate', '$200 to $600 for the first 6 months')
add_kv('Final form', 'Public website at a custom domain, free tier active, paid tier wired to Stripe, predictions returned in under 30 seconds')
add_kv('Critical path', 'Data → Model → Backend → Frontend → Deployment → Launch')

# ============================================================
# TECH STACK
# ============================================================
add_heading('Tech stack decisions', level=1)

doc.add_paragraph(
    'Every choice below is optimized for solo-founder execution: minimal moving '
    'parts, well-documented, free or near-free at the start, scalable later.'
)

table_stack = doc.add_table(rows=10, cols=3)
table_stack.style = 'Light Grid'
hdr = table_stack.rows[0].cells
hdr[0].text = 'Layer'
hdr[1].text = 'Choice'
hdr[2].text = 'Why'

stack_rows = [
    ('Foundation model', 'ESM-2 (open source) via Hugging Face', 'Free, no API key needed at small scale, runs locally if needed'),
    ('Classifier on top', 'Scikit-learn (RandomForest or XGBoost)', 'Lightweight, easy to debug, transparent'),
    ('Backend API', 'FastAPI (Python)', 'Same language as the model, auto-generates docs, fast to build'),
    ('Frontend', 'Next.js with Tailwind CSS', 'Modern, well-supported, fits Vercel deployment perfectly'),
    ('Database', 'SQLite (start) → PostgreSQL (later)', 'SQLite is one file, no setup. Migrate when scale demands.'),
    ('Authentication', 'Clerk or Auth.js', 'Free at low usage, handles signup/login/passwords'),
    ('Payments', 'Stripe', 'Industry standard, no realistic alternative'),
    ('Hosting (frontend)', 'Vercel', 'Free tier, automatic SSL, custom domain support, great DX'),
    ('Hosting (backend)', 'Railway or Render', 'Cheap, simple, handles Python apps well'),
]
for i, row in enumerate(stack_rows, start=1):
    cells = table_stack.rows[i].cells
    for j, val in enumerate(row):
        cells[j].text = val

doc.add_paragraph()

# ============================================================
# THE 10 PHASES
# ============================================================
add_heading('The 10 phases — start to launch', level=1)

doc.add_paragraph(
    'Each phase has a goal, ordered steps, deliverables, time estimate, and '
    'risks with mitigations. Phases are sequential where indicated; some '
    'sub-tasks within a phase can run in parallel.'
)

# --- Phase 0 ---
add_phase(
    num=0,
    title='Setup and infrastructure',
    goal='Have a clean working environment, accounts, and project structure ready before writing any product code.',
    time='1 weekend (~6 hours)',
    steps=[
        'Create the GitHub repository for the project (private at first, public later if open-sourced).',
        'Set up local development environment: Python 3.11+, Node.js 18+, Git.',
        'Decide on the domain name. Suggested candidates: nativeready.app, nativeready.io, nativeready.bio. Check availability.',
        'Sign up for cloud accounts: Vercel (free), Railway or Render (free tier), Hugging Face (free), Stripe (free), Cloudflare (free).',
        'Create the project folder structure: data/, model/, backend/, frontend/, docs/.',
        'Write a README with the project description and goals.',
        'Initialize Git, make the first commit, push to GitHub.',
    ],
    deliverables=[
        'Working Python and Node environments on the laptop',
        'GitHub repo with initial structure',
        'Domain name purchased ($10-$50 depending on TLD)',
        'All cloud accounts created',
    ],
    risks=[
        ('Wrong tech stack chosen, requires rewrite later',
         'Stick to the stack defined above. It is proven for this use case. Resist the urge to over-engineer.'),
        ('Domain name taken or expensive',
         'Have 5 backup domain names ready before checking. Avoid .ai if budget-constrained ($70+/yr).'),
        ('Account setup blocked by verification issues',
         'Use email and phone numbers already verified. Do not use throwaway accounts.'),
        ('Time overrun on setup',
         'Hard cap: if setup takes more than 8 hours, ship what you have and move on. Setup perfection is not the goal.'),
    ],
)

# --- Phase 1 ---
add_phase(
    num=1,
    title='Data collection',
    goal='Build a labeled dataset of protein sequences with native MS suitability outcomes.',
    time='2 weekends (~15 hours)',
    steps=[
        'Identify ~200 positive examples: proteins published in native MS papers with confirmed usable spectra. Search PubMed for "native mass spectrometry" + protein name.',
        'Pull the protein sequences from UniProt for each positive example.',
        'Identify ~200 proxy negative examples: proteins that failed in crystallography or cryo-EM (PDB has structures with REMARK entries about difficult preparation).',
        'Build a CSV: columns = sequence, label (1 = worked / 0 = failed), source paper (DOI), notes.',
        'Add basic features for each sequence: length, molecular weight, isoelectric point, predicted disorder %, predicted glycosylation sites count.',
        'Hold out 20% of the data as a test set. Save as test_set.csv.',
        'Document the data collection process so it can be re-run as new papers come out.',
    ],
    deliverables=[
        'training_set.csv with ~320 labeled examples',
        'test_set.csv with ~80 labeled examples',
        'data_collection.md documenting how it was built',
        'A list of papers used as sources',
    ],
    risks=[
        ('Not enough labeled data for good model performance',
         'Start with 400 examples; if model is weak, expand to 1000. Quality > quantity for the first pass.'),
        ('Labeling bias (only published successes get into the positive set)',
         'Acknowledge this in the model card. Use proxy negatives from crystallography literature to balance.'),
        ('Time overrun on data collection',
         'Cap at 15 hours. If incomplete, ship with smaller dataset and improve later.'),
        ('Copyright concerns on scraping papers',
         'Only collect: sequence (in UniProt, public), DOI (public metadata), and binary outcome (a fact, not copyrightable).'),
    ],
)

# --- Phase 2 ---
add_phase(
    num=2,
    title='Model development',
    goal='Train a working classifier that predicts native MS suitability with reasonable accuracy.',
    time='2 weekends (~15 hours)',
    steps=[
        'Set up an inference pipeline: feed each sequence through ESM-2 to get a 1280-dim embedding.',
        'Combine ESM embeddings with the basic biochemical features from Phase 1 (length, pI, disorder, glycosylation).',
        'Train a baseline classifier: logistic regression or RandomForest. Measure accuracy, precision, recall, AUC on the held-out test set.',
        'If baseline is poor (AUC < 0.65), iterate: try XGBoost, try only ESM embeddings, try only biochemical features, try ensemble.',
        'Once acceptable (AUC > 0.70), save the model with joblib or pickle.',
        'Build the prediction function: takes a sequence, returns suitability score (0-100), confidence interval, and risk factor breakdown.',
        'Write 20 unit tests on known sequences to catch regressions.',
    ],
    deliverables=[
        'Trained model file (~5-50 MB)',
        'Prediction function with documented input/output',
        'Performance report (AUC, confusion matrix, examples)',
        'Test suite that runs in under 1 minute',
    ],
    risks=[
        ('Model accuracy is too low to be useful',
         'Acceptable threshold is AUC > 0.70 for v1. Below that, do not launch. Pivot to a different angle (e.g., flag only obvious failures).'),
        ('ESM-2 embeddings are too slow for live API',
         'Pre-compute embeddings for common sequences. Cache results. For new sequences, accept 5-10 second response time.'),
        ('Compute costs explode',
         'Use ESM-2 Small (8M params) instead of ESM-2 Large (650M params) at first. Accept slightly lower quality for free local inference.'),
        ('Overfitting to small dataset',
         'Use cross-validation. Hold out 20% strictly. Add regularization. Document the limitation in the model card.'),
    ],
)

# --- Phase 3 ---
add_phase(
    num=3,
    title='Backend API',
    goal='Build a working API that accepts a sequence and returns the prediction in JSON format.',
    time='1 weekend (~8 hours)',
    steps=[
        'Initialize FastAPI project. One endpoint: POST /predict.',
        'Endpoint accepts JSON: { "sequence": "MGSSHHHHH..." } and returns the full prediction (score, confidence, risk factors, recommendations).',
        'Add input validation: reject sequences > 2000 amino acids, reject non-amino-acid characters, reject empty input.',
        'Add rate limiting: 50 requests per IP per month for free tier (use slowapi library).',
        'Add basic logging (which sequences were predicted, when, by whom if logged in).',
        'Test API locally with curl and the FastAPI auto-generated docs at /docs.',
        'Containerize with Docker so it runs identically locally and in deployment.',
    ],
    deliverables=[
        'Working FastAPI service runnable locally with one command',
        'Auto-generated API documentation at /docs',
        'Docker setup that builds in under 5 minutes',
        'Postman or curl-based integration tests',
    ],
    risks=[
        ('Cold start times kill user experience',
         'Use a hosting provider with always-on tier (Railway $5/mo). Avoid serverless cold starts on free tiers.'),
        ('Rate limit abuse via VPN-rotating users',
         'Accept some abuse on free tier. If usage explodes, add Cloudflare for IP analysis. Do not over-engineer at v1.'),
        ('API breaks on weird input',
         'Strict input validation upfront. Sanitize, size-check, character-check. Return clear error messages.'),
        ('Costs explode if hosting bills compound',
         'Set hard billing caps on Railway/Render at $20/month for v1. Alert at $10.'),
    ],
)

# --- Phase 4 ---
add_phase(
    num=4,
    title='Frontend (the website)',
    goal='Build the public-facing website where users paste sequences and see predictions.',
    time='2 weekends (~15 hours)',
    steps=[
        'Initialize Next.js project with Tailwind CSS.',
        'Build the single landing page: headline, brief description, large text input box, predict button.',
        'Build the results display: suitability score with progress bar, color-coded risk indicator, risk factor breakdown, recommendations, confidence interval.',
        'Style to match the visual design from the NativeReady concept doc (dark theme, cyan accent, similar to the Traversa website aesthetic).',
        'Add loading state during the 5-30 second prediction wait.',
        'Add error handling: rate limit hit, sequence invalid, server error.',
        'Add a "Download PDF report" button (use jsPDF or similar).',
        'Add a "Compare another sequence" button.',
        'Make it mobile-responsive. Test on a real phone.',
        'Add basic SEO: title tag, description, Open Graph image.',
    ],
    deliverables=[
        'Working website that runs locally',
        'Single-page UI with all critical states (loading, success, error)',
        'PDF export functionality',
        'Mobile-responsive design tested on real device',
    ],
    risks=[
        ('UI looks amateur',
         'Copy the aesthetic of inceptionlabs.ai or your own Traversa site. Do not invent. Use Tailwind UI templates if needed.'),
        ('Confusing UX',
         'Show the website to 3 friends before launch. If any of them are confused about what to do, redesign.'),
        ('Mobile broken',
         'Test on phone DURING dev, not after. Use Chrome DevTools mobile preview every commit.'),
        ('Slow page load',
         'Use Vercel CDN. Optimize images. Lazy-load anything not needed for first paint.'),
    ],
)

# --- Phase 5 ---
add_phase(
    num=5,
    title='Authentication and payments',
    goal='Allow users to sign up, upgrade to Pro, and pay $49/month via Stripe.',
    time='1 weekend (~8 hours)',
    steps=[
        'Add Clerk or Auth.js for user signup/login. Keep it simple: email + password, plus Google sign-in.',
        'Set up Stripe account, create the Pro product ($49/month subscription).',
        'Wire up Stripe Checkout: clicking "Upgrade to Pro" sends user to a Stripe-hosted checkout page.',
        'Handle the Stripe webhook for subscription events (created, canceled, payment failed).',
        'Add a usage tracker: count predictions per user per month, enforce limits.',
        'Add a billing dashboard where Pro users can manage their subscription.',
        'Test the full flow: sign up → use free tier → upgrade to Pro → cancel → go back to free tier.',
    ],
    deliverables=[
        'Working signup and login',
        'Stripe Pro subscription wired and testable',
        'Usage limits enforced at the API level',
        'Email confirmation on signup and on subscription change',
    ],
    risks=[
        ('Payment integration broken',
         'Use Stripe\'s official starter templates. Test in Stripe test mode extensively before going live.'),
        ('Subscription bugs charge users twice',
         'Start with a single Pro tier. Add Lab/Enterprise later. Less complexity = fewer bugs.'),
        ('User can\'t cancel and gets angry',
         'Use Stripe customer portal (built-in). Self-serve cancel. Refund on request, no questions, for v1.'),
        ('Authentication compromised',
         'Use Clerk or Auth.js. Do not roll your own auth. Period.'),
    ],
)

# --- Phase 6 ---
add_phase(
    num=6,
    title='Testing and quality',
    goal='Make sure nothing obvious breaks before public launch.',
    time='1 weekend (~6 hours)',
    steps=[
        'Run model on 50 known sequences (mix of positives and negatives). Verify predictions match expected outcomes within reasonable tolerance.',
        'Test edge cases: very short sequences (< 50 aa), very long (> 1500 aa), unusual amino acids (X, B, Z), empty input, special characters.',
        'Test the full payment flow end-to-end with a real test card.',
        'Test on Chrome, Safari, Firefox, mobile Safari, mobile Chrome.',
        'Test load: hit the API with 100 concurrent requests, watch for errors.',
        'Have 3 friends try the website without instructions. Watch where they get stuck.',
        'Fix the top 5 issues found. Document the rest as known issues.',
    ],
    deliverables=[
        'Test report with pass/fail on each scenario',
        'List of known issues prioritized by severity',
        'User feedback notes from 3 friend testers',
    ],
    risks=[
        ('Critical bug found in production',
         'Fix-forward fast. Have a single-command rollback ready (Vercel makes this trivial).'),
        ('Quality issues found by experts after launch',
         'Be transparent: include a "model limitations" page. Acknowledge what the tool does well and where it falls short.'),
        ('Browser-specific bugs',
         'Test the top 3 browsers minimum. Accept Safari quirks as a fact of life.'),
        ('Performance under load',
         'Vercel auto-scales. The bottleneck will be the backend API. Set up monitoring; scale Railway plan if needed.'),
    ],
)

# --- Phase 7 ---
add_phase(
    num=7,
    title='Deployment',
    goal='Get the website live at the custom domain with HTTPS.',
    time='1 weekend (~6 hours)',
    steps=[
        'Push the frontend to Vercel. Connect to the GitHub repo for auto-deploys.',
        'Push the backend to Railway. Connect to the GitHub repo for auto-deploys.',
        'Set up environment variables: Stripe keys, database URL, model paths.',
        'Configure custom domain on Vercel (point DNS at Vercel; SSL is automatic).',
        'Configure CORS so the frontend can call the backend.',
        'Verify the live site works end-to-end: predict, sign up, pay.',
        'Set up monitoring: Vercel analytics (free), Sentry for error tracking (free tier), simple uptime check.',
        'Set up automated backups of the production database.',
    ],
    deliverables=[
        'Live website at the custom domain (https://nativeready.app or similar)',
        'Auto-deploy pipeline from GitHub',
        'Monitoring dashboards',
        'Backup system in place',
    ],
    risks=[
        ('DNS or SSL configuration breaks',
         'Vercel handles both automatically if you use their nameservers. If using Cloudflare, follow their Vercel guide exactly.'),
        ('Production database lost',
         'Railway has automatic backups on paid plans. Use them. Also keep weekly manual exports.'),
        ('Costs higher than expected',
         'Set billing alerts at $20, $50, $100 across all services. Vercel and Railway both support this.'),
        ('Deployment fails repeatedly',
         'Deploy frequently in small chunks during development. Big-bang first deployment always breaks; small-step deployment is fine.'),
    ],
)

# --- Phase 8 ---
add_phase(
    num=8,
    title='Launch preparation',
    goal='Have all the materials and channels ready to make the public launch land well.',
    time='1 weekend (~6 hours)',
    steps=[
        'Polish the landing page copy. Use simple, plain language. Aim for "anyone can understand" rather than "experts will be impressed."',
        'Record a 60-second demo video using Loom. Show the workflow: paste sequence, get prediction, scan results.',
        'Write the launch announcement post for Twitter/X (3 versions: short, medium, with screenshots).',
        'Write the launch announcement for LinkedIn (longer, more professional tone).',
        'Identify 5-10 native MS researchers, students, or community accounts to mention or share with at launch.',
        'Prepare a list of subreddits to post in: r/Biochemistry, r/labrats, r/MachineLearning (with disclosure).',
        'Set up a simple analytics dashboard to watch traffic on launch day (Plausible or Vercel built-in).',
        'Prepare a feedback form (Google Form or Typeform) for early users.',
    ],
    deliverables=[
        'Polished landing page',
        '60-second Loom demo video',
        'Launch posts drafted in 3 channels (Twitter, LinkedIn, Reddit)',
        'Analytics and feedback collection ready',
    ],
    risks=[
        ('Bad copy that does not convert',
         'Copy the structure of successful B2B SaaS landing pages (Linear, Vercel, Stripe). Do not invent.'),
        ('No traffic on launch day',
         'Have 5+ outreach DMs ready to send manually on launch day to seed initial visitors.'),
        ('Demo video too long or too polished',
         'Aim for 60 seconds, raw Loom recording, no editing. Authenticity beats polish for a soft launch.'),
        ('Launch day overwhelms you',
         'Schedule launch on a day you have no other obligations. Block calendar for monitoring and replying.'),
    ],
)

# --- Phase 9 ---
add_phase(
    num=9,
    title='Soft launch',
    goal='Put the product in front of real users in the native MS community and observe what happens.',
    time='1 weekend + 1 week of monitoring (~10 hours)',
    steps=[
        'Post the launch announcement on Twitter/X. Tag relevant accounts.',
        'Post on LinkedIn. Share to your network.',
        'Post on r/Biochemistry and r/labrats with a clear disclaimer that you built this and want feedback.',
        'DM 5-10 specific researchers you have identified. Personal, short messages: "I built this thing, would love your honest 2-minute review."',
        'Monitor analytics, error logs, and social media for 7 days.',
        'Respond personally to every comment, message, and email within 24 hours.',
        'Track: visitor count, prediction count, signup count, payment count, errors.',
        'Note all feedback in a single doc for later prioritization.',
    ],
    deliverables=[
        'Launched product visible to the public',
        '24-hour response time on all feedback',
        'A feedback log documenting every comment received',
        'A 7-day metrics report',
    ],
    risks=[
        ('Crickets — no traffic, no signups, no feedback',
         'Most launches are crickets. Plan for a 30-day push, not a 1-day launch. Reach out individually to 50+ people in the field.'),
        ('Negative or hostile feedback from experts',
         'Respond gracefully. Most criticism is correct. Note it, fix what you can, acknowledge what you cannot.'),
        ('Server crashes from unexpected traffic',
         'Vercel auto-scales the frontend. If the backend struggles, scale Railway plan. Keep a "we are at capacity" fallback page ready.'),
        ('Bug discovered that requires immediate fix',
         'Have a clear hotfix process: branch, fix, test locally, push, auto-deploy. Should take under an hour.'),
    ],
)

# --- Phase 10 ---
add_phase(
    num=10,
    title='Iteration and decision',
    goal='Use the first 30-90 days of usage data to decide whether to invest more in NativeReady or move on.',
    time='Ongoing (~5 hours per week)',
    steps=[
        'Each week, review metrics: visitors, free signups, paid conversions, churn, errors.',
        'Each week, review the feedback log. Pick the top 1-2 improvements to ship.',
        'Ship improvements weekly. Small, frequent updates beat large infrequent ones.',
        'Collect new labeled data from users who report whether their experiments worked. Use this to retrain the model quarterly.',
        'At day 30: decide if traction is real (signal: at least 1 paying customer + at least 100 free users).',
        'At day 90: decide if traction is sustainable (signal: at least 5 paying customers + at least 500 free users + recurring use, not just one-time).',
        'At day 180: if sustained, commit to NativeReady as a real product line. If not, document learnings, archive, move on.',
    ],
    deliverables=[
        'Weekly metrics report',
        'Weekly product update (small but visible)',
        'Quarterly model retraining',
        'Clear go/no-go decision at day 30, 90, 180',
    ],
    risks=[
        ('No traction by day 90',
         'This is the most likely outcome. Have a graceful sunset plan: open-source the code, publish a methods paper, archive the website.'),
        ('Too much demand',
         'Raise prices. Add waitlist for Lab/Enterprise tiers. Hire help only when paying customers justify it.'),
        ('Vendor competition (Waters/Bruker bundles a similar tool)',
         'Differentiate on speed, accuracy, and breadth (multi-modality predictions). Move faster than they can.'),
        ('Burn out',
         'Cap weekly time at 5 hours during iteration phase. NativeReady is secondary to Traversa; do not let it eat the primary.'),
    ],
)

# ============================================================
# CRITICAL PATH AND DEPENDENCIES
# ============================================================
add_heading('Critical path and dependencies', level=1)

doc.add_paragraph(
    'Some phases must be sequential. Others can run in parallel. Here is the '
    'order with dependencies marked.'
)

dep_table = doc.add_table(rows=12, cols=3)
dep_table.style = 'Light Grid'
hdr_d = dep_table.rows[0].cells
hdr_d[0].text = 'Phase'
hdr_d[1].text = 'Depends on'
hdr_d[2].text = 'Can run parallel with'

dep_rows = [
    ('0. Setup', '—', '—'),
    ('1. Data collection', 'Phase 0', 'Phase 4 setup'),
    ('2. Model development', 'Phase 1', 'Phase 4 (the frontend can be built with mock predictions)'),
    ('3. Backend API', 'Phase 2', '—'),
    ('4. Frontend', 'Phase 0', 'Phase 1, 2, 3 (use mock data)'),
    ('5. Auth and payments', 'Phase 3, Phase 4', '—'),
    ('6. Testing', 'Phase 5', '—'),
    ('7. Deployment', 'Phase 6', '—'),
    ('8. Launch prep', 'Phase 7', 'Phase 6'),
    ('9. Soft launch', 'Phase 7, Phase 8', '—'),
    ('10. Iteration', 'Phase 9', 'Always-on'),
]
for i, row in enumerate(dep_rows, start=1):
    cells = dep_table.rows[i].cells
    for j, val in enumerate(row):
        cells[j].text = val

# ============================================================
# TOTAL COST BREAKDOWN
# ============================================================
add_heading('Total cost breakdown — first 6 months', level=1)

cost_table = doc.add_table(rows=11, cols=3)
cost_table.style = 'Light Grid'
hdr_c = cost_table.rows[0].cells
hdr_c[0].text = 'Service'
hdr_c[1].text = 'Cost'
hdr_c[2].text = 'When you pay'

cost_rows = [
    ('Domain name (.app or .io)', '$15-50/year', 'Phase 0, one-time annual'),
    ('Vercel (frontend hosting)', '$0 free tier', 'Free up to ~100K visitors/month'),
    ('Railway or Render (backend)', '$5-20/month', 'After Phase 7 deployment'),
    ('Hugging Face inference (if used)', '$0-50/month', 'Variable based on usage'),
    ('Stripe', '0% + 2.9% + $0.30 per transaction', 'Only on actual revenue'),
    ('Clerk or Auth.js', '$0 free tier', 'Free up to 10K users'),
    ('Sentry (error monitoring)', '$0 free tier', 'Free for small projects'),
    ('Plausible or analytics', '$0-9/month', 'Optional'),
    ('Cloudflare (DNS/CDN)', '$0', 'Always free for basics'),
    ('Total estimated', '$200-600 over 6 months', '—'),
]
for i, row in enumerate(cost_rows, start=1):
    cells = cost_table.rows[i].cells
    for j, val in enumerate(row):
        cells[j].text = val

# ============================================================
# DECISION GATES
# ============================================================
add_heading('Decision gates — when to stop', level=1)

doc.add_paragraph(
    'Three checkpoints where you decide whether to keep going.'
)

gates = [
    ('Gate 1: After Phase 2 (model development)',
        'If model AUC is below 0.65 on held-out data, the technical thesis does not work yet. Either expand the dataset (more time) or pause and revisit later. Do NOT launch a tool that does not work.'),
    ('Gate 2: After Phase 7 (deployment, before launch)',
        'If 3 friend testers cannot understand or use the website without instructions, fix the UX before launching. A confusing soft launch is wasted.'),
    ('Gate 3: Day 90 after launch',
        'If you have fewer than 5 paying customers AND fewer than 500 free users AND no recurring usage, the product has no traction. Decide: invest more (riskier), pivot positioning (medium effort), or sunset gracefully (cleanest).'),
]
for label, text in gates:
    add_kv(label, text)

# ============================================================
# SUCCESS METRICS
# ============================================================
add_heading('Success metrics by milestone', level=1)

metric_table = doc.add_table(rows=6, cols=3)
metric_table.style = 'Light Grid'
hdr_m = metric_table.rows[0].cells
hdr_m[0].text = 'Milestone'
hdr_m[1].text = 'Free users'
hdr_m[2].text = 'Paying users'

metric_rows = [
    ('Launch day', '10-50 visitors', '0'),
    ('Day 30', '100-500', '1-3'),
    ('Day 90', '500-2000', '5-15'),
    ('Day 180', '2000-5000', '15-40'),
    ('Day 365', '5000-15000', '40-100'),
]
for i, row in enumerate(metric_rows, start=1):
    cells = metric_table.rows[i].cells
    for j, val in enumerate(row):
        cells[j].text = val

doc.add_paragraph()
doc.add_paragraph(
    'These are realistic targets, not aspirational. Most PLG B2B SaaS '
    'products take 12-18 months to hit meaningful revenue. NativeReady is '
    'no exception.'
)

# ============================================================
# CLOSING
# ============================================================
doc.add_page_break()
add_heading('How to use this document', level=1)

doc.add_paragraph(
    'Treat this as a checklist, not a contract. Do the phases in order. '
    'Skip steps if they are obvious or already done. Re-prioritize when you '
    'learn something new. The goal is a working, public, payment-enabled '
    'website that anyone can visit and use.'
)

doc.add_paragraph(
    'When you finish a phase, mark it done in your own tracking. When you '
    'change a step, update this document so the build plan stays the source '
    'of truth. The next person to read this (which might be you in 3 '
    'months) should be able to follow the plan without ambiguity.'
)

add_kv('Folder for all NativeReady artifacts',
    '/Users/bfentaw2/startup/nativeready/')
add_kv('This plan',
    '/Users/bfentaw2/startup/nativeready/build_plan.docx')
add_kv('Generator script',
    '/Users/bfentaw2/startup/nativeready/make_build_plan.py')

doc.save('/Users/bfentaw2/startup/nativeready/build_plan.docx')
print("Saved: /Users/bfentaw2/startup/nativeready/build_plan.docx")
